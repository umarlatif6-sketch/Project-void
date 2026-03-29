"""
PROJECT VOID — Brand Legal Document Generator
Generates Mutual NDA and Research Collaboration Agreement as PDF and DOCX.
"""
import io
import os

GOLD = (201, 168, 76)
WHITE = (228, 228, 235)
MUTED = (136, 136, 152)
DARK = (9, 9, 11)


def _nda_paragraphs():
    return [
        ("heading", "Mutual Non-Disclosure Agreement"),
        ("subheading", "PROJECT VOID"),
        ("body",
         "This Mutual Non-Disclosure Agreement (\"Agreement\") is entered into as of "
         "[DATE] by and between:\n\n"
         "Party A: PROJECT VOID, operated by [FOUNDER FULL LEGAL NAME], located at "
         "[ADDRESS OR JURISDICTION] (\"VOID\");\n\n"
         "Party B: [COUNTERPARTY FULL LEGAL NAME OR ENTITY NAME], located at "
         "[COUNTERPARTY ADDRESS] (\"Counterparty\").\n\n"
         "VOID and Counterparty are each a \"Party\" and together the \"Parties.\""),
        ("clause_heading", "1. Purpose"),
        ("body",
         "§ 1.1 — The Parties intend to explore a potential research collaboration, "
         "technical partnership, or investment relationship relating to PROJECT VOID's "
         "technology ecosystem, including but not limited to: acoustic steganography, "
         "Al-Jabr logic language, Adriana SCL, MycoVOID biocomputing, GriDul mesh "
         "protocol, QiSync biomechanical intelligence, and associated token economy "
         "(VTX / PEACE coin) (collectively, the \"Purpose\")."),
        ("clause_heading", "2. Confidential Information"),
        ("body",
         "§ 2.1 Definition — \"Confidential Information\" means any non-public information "
         "disclosed by one Party (\"Disclosing Party\") to the other (\"Receiving Party\") "
         "in connection with the Purpose, whether disclosed orally, in writing, "
         "electronically, or by any other means, and whether or not marked \"confidential.\" "
         "This includes, without limitation: source code, technical designs, research "
         "findings, prior art filings, prototype specifications, token economic models, "
         "node architecture, and business plans.\n\n"
         "§ 2.2 Exclusions — Confidential Information does not include information that: "
         "(a) is or becomes publicly available through no breach of this Agreement; "
         "(b) was already known to the Receiving Party prior to disclosure, as evidenced "
         "by written records; (c) is independently developed by the Receiving Party without "
         "use of Confidential Information; or (d) is rightfully received from a third party "
         "without restriction."),
        ("clause_heading", "3. Obligations"),
        ("body",
         "§ 3.1 — Each Receiving Party shall: (a) hold all Confidential Information in "
         "strict confidence; (b) not disclose Confidential Information to any third party "
         "without prior written consent of the Disclosing Party; (c) use Confidential "
         "Information solely for the Purpose; (d) limit access to Confidential Information "
         "to those representatives who have a need to know and are bound by confidentiality "
         "obligations no less restrictive than those in this Agreement.\n\n"
         "§ 3.2 Required Disclosure — If required by law or court order to disclose "
         "Confidential Information, the Receiving Party shall provide prompt prior written "
         "notice to the Disclosing Party (where legally permissible) and cooperate with "
         "the Disclosing Party's efforts to seek a protective order."),
        ("clause_heading", "4. Intellectual Property"),
        ("body",
         "§ 4.1 — Nothing in this Agreement grants either Party any right, title, or "
         "interest in the other Party's Confidential Information or intellectual property. "
         "All prior art filings, inventions, and proprietary systems disclosed by VOID "
         "remain the exclusive property of VOID."),
        ("clause_heading", "5. Term and Termination"),
        ("body",
         "§ 5.1 — This Agreement commences on the date first written above and continues "
         "until either Party terminates it with 30 days' written notice. Confidentiality "
         "obligations survive termination and continue for a period of two (2) years from "
         "the date of termination."),
        ("clause_heading", "6. Governing Law & Dispute Resolution"),
        ("body",
         "§ 6.1 — This Agreement shall be governed by the laws of [STATE / COUNTRY]. "
         "Any dispute arising from this Agreement shall first be addressed through "
         "good-faith negotiation. If unresolved within 30 days, the Parties agree to "
         "binding arbitration in [CITY, STATE / COUNTRY] under the rules of "
         "[ARBITRATION BODY, e.g., AAA / JAMS / ICC]."),
        ("clause_heading", "7. General"),
        ("body",
         "§ 7.1 Entire Agreement — This Agreement constitutes the entire agreement "
         "between the Parties with respect to its subject matter and supersedes all "
         "prior discussions and understandings.\n\n"
         "§ 7.2 Amendment — No amendment to this Agreement shall be valid unless made "
         "in writing and signed by both Parties.\n\n"
         "§ 7.3 Severability — If any provision is found unenforceable, the remaining "
         "provisions remain in full force."),
        ("clause_heading", "Signatures"),
        ("body",
         "By signing below, each Party agrees to be bound by the terms of this Agreement.\n\n"
         "Party A — PROJECT VOID\n"
         "Signature: _____________________________\n"
         "Name: [FOUNDER FULL NAME]\n"
         "Date: _____________________________\n\n"
         "Party B — Counterparty\n"
         "Signature: _____________________________\n"
         "Name: [COUNTERPARTY NAME]\n"
         "Date: _____________________________"),
        ("footer", "PROJECT VOID — Mutual NDA | Template. Complete bracketed fields before signing. Not legal advice."),
    ]


def _rca_paragraphs():
    return [
        ("heading", "Research Collaboration Agreement"),
        ("subheading", "PROJECT VOID"),
        ("body",
         "This Research Collaboration Agreement (\"Agreement\") is entered into as of "
         "[DATE] by and between:\n\n"
         "VOID: PROJECT VOID, operated by [FOUNDER FULL LEGAL NAME], located at "
         "[ADDRESS];\n\n"
         "Collaborator: [INSTITUTION OR RESEARCHER NAME], located at [ADDRESS]."),
        ("clause_heading", "1. Research Scope"),
        ("body",
         "§ 1.1 Project — The Parties agree to collaborate on the following research "
         "project: [PROJECT TITLE — e.g., \"Mycelium Network Signal Propagation in "
         "MRB-4000 Architecture\" or \"GriDul Mesh Consensus Performance Under Adversarial "
         "Conditions\" or \"QiSync Biorhythm-Machine State Synchronisation Protocols\"] "
         "(the \"Project\").\n\n"
         "§ 1.2 Objectives — The Parties will work together to achieve the following "
         "objectives: [LIST 2–4 SPECIFIC RESEARCH OBJECTIVES].\n\n"
         "§ 1.3 Duration — The Project shall commence on [START DATE] and conclude on "
         "[END DATE], unless extended by mutual written agreement."),
        ("clause_heading", "2. Responsibilities"),
        ("body",
         "§ 2.1 VOID Responsibilities — VOID shall: (a) provide access to the relevant "
         "technical systems and documentation as specified in Schedule A; (b) contribute "
         "[X hours/week or specific deliverables] of engineering time; (c) designate "
         "[FOUNDER NAME] as the primary point of contact.\n\n"
         "§ 2.2 Collaborator Responsibilities — Collaborator shall: (a) provide "
         "[SPECIFIC CONTRIBUTION — e.g., laboratory facilities, domain expertise, "
         "computing resources]; (b) designate [COLLABORATOR CONTACT NAME] as the "
         "primary point of contact."),
        ("clause_heading", "3. Intellectual Property"),
        ("body",
         "§ 3.1 Background IP — \"Background IP\" means intellectual property developed "
         "independently by a Party prior to or outside this Agreement. Each Party retains "
         "exclusive ownership of its Background IP. Background IP of VOID includes all "
         "prior art filings, existing technical architecture, the Al-Jabr/Adriana language "
         "system, the acoustic steganography engine, and all related inventions documented "
         "prior to the commencement date of this Agreement.\n\n"
         "§ 3.2 Foreground IP — \"Foreground IP\" means intellectual property created "
         "jointly by both Parties in the course of the Project. Foreground IP shall be "
         "co-owned equally by both Parties. Each Party may use Foreground IP for "
         "non-commercial research purposes without the other's consent. Commercial use "
         "of Foreground IP requires prior written agreement on terms, including revenue "
         "sharing.\n\n"
         "§ 3.3 Solo Inventions — Intellectual property created solely by one Party in "
         "the course of the Project, without material contribution from the other, shall "
         "be owned by the creating Party. The creating Party shall notify the other Party "
         "within 30 days of identifying a solo invention."),
        ("clause_heading", "4. Publication & Confidentiality"),
        ("body",
         "§ 4.1 Publication Review — Either Party wishing to publish results arising "
         "from the Project shall provide the other Party with a draft manuscript at least "
         "60 days before submission. The reviewing Party may request a delay of up to "
         "90 additional days solely for the purpose of filing patent applications on "
         "identified patentable inventions.\n\n"
         "§ 4.2 Attribution — All publications arising from the Project shall appropriately "
         "credit both Parties. VOID's contributions shall be credited as \"PROJECT VOID\" "
         "with authorship attributed to [FOUNDER NAME].\n\n"
         "§ 4.3 Confidentiality — The Parties' confidentiality obligations are governed "
         "by a separate Mutual NDA or, if no separate NDA exists, by the confidentiality "
         "provisions set out in Schedule B to this Agreement."),
        ("clause_heading", "5. Resources & Funding"),
        ("body",
         "§ 5.1 — Each Party shall bear its own costs unless otherwise specified in "
         "Schedule C. Any shared costs shall be allocated as follows: [DESCRIBE COST "
         "SHARING]. No party is obligated to fund the other's operations."),
        ("clause_heading", "6. Term, Termination & Consequences"),
        ("body",
         "§ 6.1 Term — This Agreement commences on the date first written above and "
         "continues for the duration of the Project unless terminated earlier.\n\n"
         "§ 6.2 Termination for Convenience — Either Party may terminate this Agreement "
         "with 30 days' written notice. Upon termination, each Party retains rights to "
         "Background IP. Ownership of Foreground IP created before termination is "
         "unaffected.\n\n"
         "§ 6.3 Termination for Cause — Either Party may terminate immediately upon "
         "written notice if the other Party materially breaches this Agreement and fails "
         "to cure such breach within 15 days of written notice."),
        ("clause_heading", "7. Representations & Warranties"),
        ("body",
         "§ 7.1 — Each Party represents that: (a) it has the authority to enter into "
         "this Agreement; (b) entering into this Agreement does not violate any other "
         "agreement to which it is a party; (c) it will conduct all research activities "
         "in compliance with applicable laws and ethical standards."),
        ("clause_heading", "8. Governing Law & Dispute Resolution"),
        ("body",
         "§ 8.1 — This Agreement shall be governed by the laws of [STATE / COUNTRY]. "
         "Disputes shall be resolved first by good-faith negotiation, then by binding "
         "arbitration in [CITY, STATE / COUNTRY] under the rules of [ARBITRATION BODY]."),
        ("clause_heading", "9. Schedules"),
        ("body",
         "Schedule A — Technical Access: [List specific systems, codebases, or "
         "documentation provided by VOID]\n\n"
         "Schedule B — Confidentiality Terms (if no separate NDA): [Attach or "
         "incorporate NDA terms]\n\n"
         "Schedule C — Cost Allocation: [Itemise shared costs and responsibilities]"),
        ("clause_heading", "Signatures"),
        ("body",
         "By signing below, each Party agrees to be bound by the terms of this Agreement.\n\n"
         "PROJECT VOID\n"
         "Signature: _____________________________\n"
         "Name: [FOUNDER FULL NAME]\n"
         "Date: _____________________________\n\n"
         "Collaborator\n"
         "Signature: _____________________________\n"
         "Name & Title: [NAME, TITLE]\n"
         "Date: _____________________________"),
        ("footer", "PROJECT VOID — Research Collaboration Agreement | Template. Complete bracketed fields before signing. Not legal advice."),
    ]


def _ascii_safe(text: str) -> str:
    """Replace common Unicode typographic characters with ASCII equivalents."""
    replacements = {
        "\u2014": "--",   # em dash
        "\u2013": "-",    # en dash
        "\u2018": "'",    # left single quote
        "\u2019": "'",    # right single quote
        "\u201c": '"',    # left double quote
        "\u201d": '"',    # right double quote
        "\u2026": "...",  # ellipsis
        "\u00a0": " ",    # non-breaking space
        "\u2022": "*",    # bullet
        "\u2192": "->",   # right arrow
        "\u00e9": "e",    # e-acute
        "\u00e0": "a",    # a-grave
        "\u00e8": "e",    # e-grave
        "\u00fc": "u",    # u-umlaut
        "\u00e4": "a",    # a-umlaut
        "\u00f6": "o",    # o-umlaut
        "\u00df": "ss",   # sharp-s
        "\u00b0": " deg", # degree sign
        "\u00d7": "x",    # multiplication sign
        "\u2212": "-",    # minus sign
        "\u00ab": "<<",   # left guillemet
        "\u00bb": ">>",   # right guillemet
    }
    for uni, asc in replacements.items():
        text = text.replace(uni, asc)
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text


def generate_pdf(doc_type: str) -> bytes:
    """Generate a PDF for 'nda' or 'rca'. Returns bytes."""
    from fpdf import FPDF

    paragraphs = _nda_paragraphs() if doc_type == "nda" else _rca_paragraphs()

    pdf = FPDF()
    pdf.set_margins(25, 25, 25)
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()
    pdf.set_fill_color(9, 9, 11)

    for kind, text in paragraphs:
        safe = _ascii_safe(text)
        if kind == "heading":
            pdf.set_font("Helvetica", "B", 18)
            pdf.set_text_color(201, 168, 76)
            pdf.multi_cell(0, 10, safe, ln=True)
            pdf.ln(2)
        elif kind == "subheading":
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(136, 136, 152)
            pdf.multi_cell(0, 7, safe, ln=True)
            pdf.ln(4)
        elif kind == "clause_heading":
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(201, 168, 76)
            pdf.multi_cell(0, 7, safe, ln=True)
            pdf.ln(1)
        elif kind == "body":
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(200, 200, 210)
            pdf.multi_cell(0, 6, safe, ln=True)
            pdf.ln(3)
        elif kind == "footer":
            pdf.set_font("Helvetica", "I", 8)
            pdf.set_text_color(100, 100, 110)
            pdf.ln(6)
            pdf.set_draw_color(50, 50, 60)
            pdf.line(25, pdf.get_y(), 185, pdf.get_y())
            pdf.ln(3)
            pdf.multi_cell(0, 5, safe, ln=True)

    return pdf.output()


def generate_docx(doc_type: str) -> bytes:
    """Generate a DOCX for 'nda' or 'rca'. Returns bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraphs = _nda_paragraphs() if doc_type == "nda" else _rca_paragraphs()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for sect in doc.sections:
        sect.top_margin = Inches(1)
        sect.bottom_margin = Inches(1)
        sect.left_margin = Inches(1.25)
        sect.right_margin = Inches(1.25)

    for kind, text in paragraphs:
        if kind == "heading":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(201, 168, 76)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(4)
        elif kind == "subheading":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(136, 136, 152)
            p.paragraph_format.space_after = Pt(8)
        elif kind == "clause_heading":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(201, 168, 76)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        elif kind == "body":
            for line in text.split("\n\n"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(4)
        elif kind == "footer":
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(120, 120, 130)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
